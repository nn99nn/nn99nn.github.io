from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

GOLD = RGBColor(0xB8, 0x96, 0x46)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x88, 0x88, 0x88)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_section_header(text_zh, text_en, text_ja):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    # Chinese
    r = p.add_run(text_zh + "　")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = GOLD
    # English
    r2 = p.add_run(text_en + "　")
    r2.font.size = Pt(10)
    r2.font.bold = False
    r2.font.color.rgb = GOLD
    # Japanese
    r3 = p.add_run(text_ja)
    r3.font.size = Pt(9)
    r3.font.color.rgb = GOLD
    # Divider
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    r4 = p2.add_run("─" * 40)
    r4.font.size = Pt(8)
    r4.font.color.rgb = GOLD

def add_dish(zh, en, ja, note_zh="", note_en="", note_ja="", price="", indent=False):
    # Chinese name
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(zh)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = BLACK
    if price:
        r2 = p.add_run(f"　{price}")
        r2.font.size = Pt(11)
        r2.font.bold = False
        r2.font.color.rgb = GRAY

    # English
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    if indent:
        p2.paragraph_format.left_indent = Cm(0.5)
    r3 = p2.add_run(en)
    r3.font.size = Pt(10)
    r3.font.color.rgb = BLACK

    # Japanese
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(1)
    if indent:
        p3.paragraph_format.left_indent = Cm(0.5)
    r4 = p3.add_run(ja)
    r4.font.size = Pt(9)
    r4.font.color.rgb = GRAY

    # Notes
    if note_zh or note_en or note_ja:
        p4 = doc.add_paragraph()
        p4.paragraph_format.space_before = Pt(0)
        p4.paragraph_format.space_after = Pt(2)
        if indent:
            p4.paragraph_format.left_indent = Cm(0.5)
        if note_zh:
            rn1 = p4.add_run(f"({note_zh})")
            rn1.font.size = Pt(8.5)
            rn1.font.color.rgb = GRAY
        if note_en:
            rn2 = p4.add_run(f"  ({note_en})")
            rn2.font.size = Pt(8.5)
            rn2.font.color.rgb = GRAY
        if note_ja:
            rn3 = p4.add_run(f"  ({note_ja})")
            rn3.font.size = Pt(8.5)
            rn3.font.color.rgb = GRAY

# ─── Title ────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
rt = title.add_run("捌伍　菜單")
rt.font.size = Pt(22)
rt.font.bold = True
rt.font.color.rgb = GOLD

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(16)
rs = sub.add_run("85TD Menu　メニュー")
rs.font.size = Pt(11)
rs.font.color.rgb = GRAY

# ─── 燒味 ─────────────────────────────────────────────────
add_section_header("燒味", "BBQ & Roasted Meats", "焼き物")

add_dish(
    "潮式滷水鵝拼盤",
    "Teochew-Style Braised Goose Platter",
    "潮州風ガチョウのブレイズ盛り合わせ",
    "豆腐・花生・鵝肉・鵝肝・鵝翅",
    "Tofu · Peanuts · Goose · Goose Liver · Wing",
    "豆腐・落花生・ガチョウ肉・レバー・手羽",
    "850"
)
add_dish(
    "添第靚燒鴨 / 鵝",
    "85TD Signature Roasted Duck / Goose",
    "85TD 特製ロースト鴨・ガチョウ",
)
add_dish(
    "燒味拼盤",
    "Assorted BBQ Platter",
    "焼き物盛り合わせ",
    "二款 / 三款",
    "2 Items / 3 Items",
    "2種 / 3種",
    "1480 / 1580"
)
add_dish(
    "捌伍叉燒皇",
    "85 Signature Premium Char Siu",
    "85 特製プレミアムチャーシュー",
    price="1380"
)
add_dish(
    "化皮燒乳豬",
    "Cantonese Crispy-Skin Roasted Suckling Pig",
    "広東風クリスピー皮の仔豚ロースト",
    "例 / 半隻 / 1隻",
    "Portion / Half / Whole",
    "1皿 / ハーフ / 1頭",
    "1580 / 5080 / 10080"
)
add_dish(
    "瑤柱貴妃雞",
    "Conpoy Superior Broth Poached Chicken",
    "干し貝柱の上湯仕立て貴妃鶏",
    "例 / 半隻 / 1隻",
    "Portion / Half / Whole",
    "1皿 / ハーフ / 1羽",
    "760 / 960 / 1880"
)

# ─── 頭盤與小食 ────────────────────────────────────────────
add_section_header("頭盤與小食", "Starters & Bites", "前菜・一品料理")

add_dish("烏梅蘿蔔", "Pickled Radish with Smoked Plum", "燻製梅風味の大根漬け", price="250")
add_dish("醬香茄子", "Braised Eggplant in Savory Bean Sauce", "醤油ソースのナス煮込み", price="280")
add_dish("鮮沙薑涷豬手", "Chilled Pig's Trotter with Fresh Sand Ginger", "砂ショウガ風味の冷製豚足ゼリー寄せ", price="520")
add_dish("拍蒜青瓜海蜇頭", "Smashed Cucumber & Jellyfish with Crushed Garlic", "たたきキュウリとクラゲのにんにく和え", price="650")
add_dish("水晶餚肉", "Crystal Pork Terrine", "豚肉の水晶煮こごり", price="580")
add_dish("川味皮蛋豆腐", "Sichuan-Style Century Egg Tofu", "四川風ピータン豆腐", price="280")
add_dish("椒麻南非凍鮮鮑", "Chilled South African Abalone with Sichuan Numbing Pepper Sauce", "南アフリカ産鮑の冷製・山椒痺れソース", price="680")
add_dish("金沙脆魚皮", "Crispy Fish Skin with Salted Egg Yolk", "塩卵黄がけサクサク魚皮揚げ", price="360")
add_dish("麻辣脆鴨血", "Spicy Crispy Duck Blood Curd", "麻辣風味サクサク鴨血揚げ", price="360")
add_dish("糖醋百花油條", "Sweet & Sour Shrimp Paste Stuffed Cruller", "エビすり身詰めの揚げパン・甘酢ソース", price="580")
add_dish(
    "海膽龍皇戈渣",
    "Sea Urchin & Lobster Bisque Crispy Savory Custard",
    "ウニとロブスタービスク風味のサクサク揚げ蒸し",
    price="680"
)
add_dish("桂花燒鱔球", "Osmanthus Sugar-Glazed Eel Balls", "キンモクセイ糖がけ鰻の揚げ団子", price="780")

# ─── 點心 ────────────────────────────────────────────────
add_section_header("點心", "Dim Sum", "点心")

add_dish("捌伍蝦餃皇（3件）", "85TD Signature Har Gow (3 pcs)", "捌伍特製海老蒸し餃子（3個）", price="390")
add_dish("茶香鵪鶉蛋燒賣（3件）", "Tea-Infused Quail Egg Siu Mai (3 pcs)", "お茶香る鶉卵入りシュウマイ（3個）", price="360")
add_dish("酸菜金魚餃（2件）", "Goldfish-Shaped Pickled Mustard Green Dumpling (2 pcs)", "金魚形酸菜餃子（2個）", price="280")
add_dish("牛肝菌素餃（3件）", "Porcini Mushroom Vegetarian Dumpling (3 pcs)", "ポルチーニ茸の精進餃子（3個）", price="270")
add_dish("蝦湯小籠包（3件）", "Shrimp Broth Xiao Long Bao (3 pcs)", "海老スープ小籠包（3個）", price="360")
add_dish("醬皇蒸鳳爪（例）", "Steamed Chicken Feet in House Signature Sauce (per portion)", "特製醤汁蒸し鶏の爪（1人前）", price="220")
add_dish("千絲蘿蔔酥（3件）", "Shredded Turnip Puff Pastry (3 pcs)", "千切り蘿蔔のパイ（3個）", price="300")
add_dish("海皇天鵝酥（3件）", "Seafood Swan Puff Pastry (3 pcs)", "海鮮白鳥形パイ（3個）", price="360")
add_dish("鮑魚雞粒酥（1件）", "Abalone & Diced Chicken Puff Pastry (1 pc)", "アワビと鶏肉のパイ（1個）", price="360")
add_dish("XO醬炒蝦米腸（例）", "Cheung Fun Stir-Fried with XO Sauce & Dried Shrimp (per portion)", "XOソース干し海老炒め腸粉（1人前）", price="360")
add_dish("脆皮海皇腸粉（例）", "Crispy Seafood Cheung Fun (per portion)", "クリスピー海鮮腸粉（1人前）", price="350")
add_dish("香菜叉燒腸粉（例）", "Cilantro & Char Siu Cheung Fun (per portion)", "パクチーと叉焼腸粉（1人前）", price="300")
add_dish("金耳素菌腸粉（例）", "Golden Ear Fungus & Mushroom Vegetarian Cheung Fun (per portion)", "金耳きのこの精進腸粉（1人前）", price="280")

# ─── 湯羹類 ────────────────────────────────────────────────
add_section_header("湯羹類", "Soups", "スープ")

add_dish(
    "魚翅花膠燴燕窩",
    "Fish Maw & Bird's Nest in Saffron Golden Broth",
    "サフラン・鴨・鶏・豚の黄金スープ・花膠と燕の巣の煮込み",
    "Duck · Chicken · Pork · Tibetan Saffron",
    "鴨・鶏・豚・西藏紅花",
    "位上"
)
add_dish("瑤柱花膠鴨絲羹", "Conpoy, Fish Maw & Shredded Duck Thick Soup", "干し貝柱・花膠・鴨肉のとろみスープ", price="位上")
add_dish("雪耳海鮮豆腐羹", "Snow Fungus, Seafood & Tofu Thick Soup", "白きくらげ・海鮮・豆腐のとろみスープ", price="位上")
add_dish("松茸瑤柱燉老雞", "Double-Boiled Aged Hen with Matsutake & Conpoy", "松茸と干し貝柱の老鶏炖スープ", price="位上")
add_dish("羊肚菌鮑魚燉鷓鴣", "Double-Boiled Partridge with Morel Mushroom & Abalone", "モリーユ茸と鮑の鷓鴣炖スープ", price="位上")
add_dish("捌伍濃湯佛跳牆", "85 Signature Buddha Jumps Over the Wall", "85 特製「仏跳牆」濃厚スープ", price="位上")
add_dish("清燉素獅子頭", "Clear-Braised Vegetarian Lion's Head", "澄んだスープで炖た精進ライオンズヘッド", price="位上")

# ─── 海鮮類 ────────────────────────────────────────────────
add_section_header("海鮮類", "Seafood", "海鮮")

add_dish("焗釀鮮蟹蓋", "Baked Crab Shell Stuffed with Crabmeat, Onion & Cream", "カニの甲羅焼き・カニ肉・玉ねぎ・クリーム詰め", price="位上")
add_dish("蒜香金菇蒸帶子", "Steamed Scallops with Garlic & Enoki Mushroom", "ニンニク香るえのき茸と帆立蒸し", price="位上")
add_dish("黑松露蛋白蒸龍蝦球", "Steamed Lobster Medallions with Black Truffle & Egg White", "黒トリュフと卵白蒸しロブスターの団子", price="位上")
add_dish("老菜脯高樑炆龍蝦尾", "Braised Lobster Tail with Aged Preserved Radish & Kaoliang Wine", "古漬け大根と金門高粱酒炆ロブスターテール", price="半隻")
add_dish("豉椒鮮露筍炒鱈魚球", "Stir-Fried Cod Balls with Asparagus, Black Bean & Chili", "アスパラガス・豆豉・唐辛子炒めタラの団子", price="例")
add_dish("上湯明蝦粉絲煲", "Superior Broth King Prawn & Glass Noodle Casserole", "上湯海老と春雨の土鍋煮込み", "2隻起", "Minimum 2 Prawns", "2尾より", "例")
add_dish("鳳凰炒帶子", "Wok-Fried Scallops in Crispy Salted Egg Yolk Phoenix Nest", "塩卵黄サクサク鳳凰巣の帆立炒め", price="例")
add_dish(
    "黑潮軟殼龍蝦",
    "Soft-Shell Lobster",
    "ソフトシェルロブスター",
    "可選：避風塘 / 金沙炒",
    "Choice of: Typhoon Shelter Style / Salted Egg Yolk Style",
    "避風塘スタイル / 塩卵黄スタイル",
    "例"
)
add_dish("桂花炒素翅", "Osmanthus Stir-Fried Vegetarian Shark's Fin", "キンモクセイ風味の精進フカヒレ炒め", price="例")
add_dish("蔥油蒸星斑球", "Steamed Leopard Grouper Medallions with Scallion Oil", "ネギ油蒸しレパードグルーパーの切り身", price="例")
add_dish("椒鹽中卷", "Salt & Pepper Squid", "塩コショウ揚げイカ", price="例")

# ─── 活海鮮 ────────────────────────────────────────────────
add_section_header("活海鮮", "Live Seafood", "活き海鮮")

add_dish(
    "清蒸或上湯龍蝦",
    "Live Lobster — Steamed or Poached in Superior Broth",
    "活きロブスター — 清蒸または上湯煮",
    "蒜蓉蒸 / 麻香剁椒蒸 / 蠔油汁炆",
    "Steamed with Minced Garlic · Steamed with Sesame & Chopped Chili · Braised in Oyster Sauce",
    "にんにく蒸し・ごま香る刻み唐辛子蒸し・牡蠣ソース炆",
    "時價"
)
add_dish(
    "南非鮮鮑魚",
    "South African Fresh Abalone",
    "南アフリカ産活き鮑",
    "蔥油鮮沙薑蒸 / 薑蔥焗 / 避風塘炒",
    "Steamed with Scallion Oil & Fresh Sand Ginger · Baked with Ginger & Scallion · Typhoon Shelter Style",
    "ネギ油・砂ショウガ蒸し・生姜ネギ焗・避風塘スタイル",
    "時價"
)
add_dish(
    "沙公 · 沙母 · 紅蟳",
    "Male Blue Swimming Crab · Female Blue Swimming Crab · Red Swimming Crab",
    "オスのガザミ · メスのガザミ · レッドクラブ",
    price="時價"
)

# ─── 家禽與肉類 ────────────────────────────────────────────
add_section_header("家禽與肉類", "Poultry & Meat", "家禽・肉料理")

add_dish(
    "茶皇燒乳鴿",
    "Tea-Fragrant Roasted Imperial Squab",
    "茶香り豊かな仔鳩のロースト",
    "二天前預訂",
    "Pre-order 2 days in advance",
    "2日前要予約"
)
add_dish(
    "富貴鮑汁八寶鴨",
    "Fortune Eight-Treasure Duck in Abalone Sauce",
    "縁起の八宝アワビソース鴨詰め",
    "乾草菇・乾鮑魚・花膠・花菇・黃肉參　三天前預訂",
    "Dried Mushroom · Dried Abalone · Fish Maw · Flower Mushroom · Sea Cucumber　Pre-order 3 days in advance",
    "乾燥きのこ・乾燥鮑・花膠・花しいたけ・なまこ　3日前要予約"
)
add_dish(
    "古法鹽焗雞",
    "Traditional Salt-Baked Chicken",
    "古法塩焗鶏",
    "文昌雞・乾蔥頭・沙薑粒・蔥段・八角　一天前預訂",
    "Wenchang Chicken · Shallot · Sand Ginger · Scallion · Star Anise　Pre-order 1 day in advance",
    "文昌鶏・エシャロット・砂ショウガ・ネギ・八角　1日前要予約"
)
add_dish(
    "捌伍脆皮雞",
    "85 Signature Crispy-Skin Chicken",
    "85 特製クリスピーチキン",
    "半隻 / 1隻",
    "Half / Whole",
    "ハーフ / 1羽"
)
add_dish("鳳梨咕嚕肉", "Sweet & Sour Pork with Pineapple", "パイナップル入り酢豚", price="例")
add_dish("手工辣腸炒海參", "Stir-Fried Sea Cucumber with Handmade Hunan Spicy Sausage", "手作り湖南辣腸と海参の炒め", price="例")
add_dish("一口陳醋骨", "Bite-Size Aged Vinegar Pork Ribs", "老酢のひと口スペアリブ", price="例")
add_dish("金耳百合炒松板肉", "Stir-Fried Pork Jowl with Golden Ear Fungus & Lily Bulb", "金耳きのこ・百合根と松板豚頬肉の炒め", price="例")
add_dish("芥末香蔥和牛粒", "Wagyu Beef Cubes with Wasabi & Scallion", "和牛サイコロステーキ・わさびとネギ風味", price="例")
add_dish(
    "藥膳牛筋腩煲",
    "Herbal Beef Tendon & Brisket Casserole",
    "薬膳牛筋・ブリスケット土鍋煮込み",
    "當歸・北芪・杞子・紅棗",
    "Angelica Root · Astragalus · Wolfberry · Red Date",
    "当帰・黄耆・クコの実・なつめ",
    "例"
)
add_dish("沙嗲溫體牛肉炒菜苗　台灣", "Satay Taiwan Fresh Beef Stir-Fried with Vegetable Shoots", "台湾産温体牛肉のサテー炒め・野菜の芽添え", price="例")
add_dish(
    "現烤日本和牛叉燒",
    "Freshly Roasted Japanese Wagyu Char Siu",
    "日本和牛の焼きたてチャーシュー",
    "一天前預訂",
    "Pre-order 1 day in advance",
    "1日前要予約"
)
add_dish(
    "青醬燒紐西蘭高地和羊排",
    "Roasted New Zealand Highland Lamb Chops with Cilantro & Cashew Sauce",
    "ニュージーランド高地産ラムチョップのパクチーとカシューナッツソース焼き"
)

# ─── 蔬菜、豆腐 ────────────────────────────────────────────
add_section_header("蔬菜、豆腐", "Vegetables & Tofu", "野菜・豆腐")

add_dish("瑤柱粉絲雜菜煲", "Conpoy, Glass Noodle & Seasonal Vegetable Casserole", "干し貝柱・春雨・雑菜の土鍋煮込み", price="例")
add_dish("金銀蛋蒜子浸時蔬", "Seasonal Vegetables Poached with Salted Egg, Century Egg & Garlic", "塩卵・ピータン・にんにくで仕上げた旬野菜のスープ浸し", price="例")
add_dish("欖菜干煸炒四季豆　素", "Dry-Fried Green Beans with Olive Vegetable (Vegetarian)", "オリーブ菜とインゲンの空炒め（精進）", price="例")
add_dish(
    "梅菜皇蒸茄子　素",
    "Steamed Eggplant with Preserved Mustard Greens in Lotus Leaf (Vegetarian)",
    "梅菜皇と茄子の蒸し物・蓮の葉包み（精進）",
    "荷葉底・大籠上",
    "Served in Lotus Leaf, Large Steamer",
    "蓮の葉包み・大蒸籠仕立て",
    "例"
)
add_dish("麻婆蝦腰豆腐煲", "Mapo Tofu Casserole with Shrimp", "海老入り麻婆豆腐土鍋", price="例")
add_dish(
    "時蔬選擇",
    "Seasonal Vegetables",
    "旬野菜",
    "蒜蓉 / 清炒 / 上湯",
    "Minced Garlic / Plain Stir-Fried / Superior Broth",
    "にんにく炒め / 清炒め / 上湯仕立て",
    "例"
)

# ─── 飯麵類 ────────────────────────────────────────────────
add_section_header("飯麵類", "Rice & Noodles", "ご飯・麺")

add_dish("XO醬炒腸粉", "Stir-Fried Rice Noodle Rolls with XO Sauce", "XOソース炒め腸粉", price="例")
add_dish("毛豆雪菜炒年糕　素", "Stir-Fried Rice Cake with Edamame & Pickled Mustard Greens (Vegetarian)", "枝豆と雪菜の炒め年糕（精進）", price="例")
add_dish("乾炒牛河", "Dry-Fried Beef Ho Fun", "乾炒め牛肉河粉", price="例")
add_dish("鳴哥炒出前一丁", "Chef Ming's Signature Stir-Fried Nissin Noodles", "鳴シェフ特製出前一丁炒め", price="例")
add_dish(
    "泡菜和牛鬆炒鴛鴦米",
    "Kimchi & Wagyu Beef Floss Yin-Yang Fried Rice",
    "キムチと和牛鬆の鴛鴦ライス炒め",
    "白米＋脆米",
    "Steamed Rice & Crispy Puffed Rice",
    "白米・揚げカリカリ米",
    "例"
)
add_dish("捌伍炒飯", "85 Signature Fried Rice", "85 特製チャーハン", price="例")
add_dish("港式豬扒焗飯", "Hong Kong-Style Baked Rice with Pork Chop", "香港風ポークチョップのオーブン焗きご飯", price="例")
add_dish("松露蛋白炒飯　素", "Truffle & Egg White Fried Rice (Vegetarian)", "トリュフと卵白の炒飯（精進）", price="例")
add_dish("鼎湖上素炆米　素", "Dinghu-Style Premium Vegetarian Braised Rice (Vegetarian)", "鼎湖山精進食材の炆ご飯（精進）", price="例")
add_dish("榨菜火鴨絲湯米", "Rice Noodle Soup with Preserved Vegetable & Roasted Duck", "搾菜と焼き鴨の米粉スープ", price="例")

# ─── 甜點 ──────────────────────────────────────────────────
add_section_header("甜點", "Dessert", "デザート")

add_dish("椰奶西米糕", "Coconut Milk Sago Pudding", "ココナッツミルクのサゴプリン")
add_dish("開心果奶凍（位上）", "Pistachio Panna Cotta (individual)", "ピスタチオのパンナコッタ（1人前）", price="220")
add_dish("楊汁甘露（位上）", "Mango Pomelo Sago (individual)", "マンゴーポメロサゴ（楊枝甘露）（1人前）", price="280")
add_dish("手磨杏仁茶湯圓（位上）", "Hand-Ground Almond Milk Tea with Glutinous Rice Balls (individual)", "手挽き杏仁茶と白玉団子（1人前）", price="260")
add_dish("流心芋頭包（3件）", "Molten Taro Bun (3 pcs)", "とろけるタロイモ包子（3個）", price="270")
add_dish("澳門蛋撻仔（3件）", "Macanese Egg Tart (3 pcs)", "マカオ風エッグタルト（3個）", price="270")
add_dish("班蘭椰汁西米糕（3件）", "Pandan Coconut Sago Cake (3 pcs)", "パンダンリーフ椰汁サゴケーキ（3個）", price="270")

# ─── 私藏 ──────────────────────────────────────────────────
add_section_header("私藏", "Hidden Menu", "秘蔵メニュー")

add_dish(
    "脆皮麻香辣子雞",
    "Crispy-Skin Sesame Spicy Chili Chicken",
    "クリスピー皮・胡麻香る麻辣チキン",
    "半隻 / 1隻",
    "Half / Whole",
    "ハーフ / 1羽"
)

output_path = "/Users/nitrowong/Documents/Claude Code/85 Menu/捌伍菜單_翻譯草稿.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
